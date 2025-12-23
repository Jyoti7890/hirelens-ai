# backend/utils/extract_text.py

import io
import os
import docx
import PyPDF2
import fitz  # PyMuPDF
import pytesseract
from PIL import Image


# ---------------- PDF HELPERS ---------------- #

def extract_pdf_pypdf2(file_path: str) -> str:
    text = ""
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                try:
                    content = page.extract_text()
                    if content:
                        text += content + "\n"
                except Exception:
                    continue
    except Exception:
        return ""
    return text.strip()


def extract_pdf_pymupdf(file_path: str) -> str:
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            try:
                content = page.get_text("text")
                if content:
                    text += content + "\n"
            except Exception:
                continue
    except Exception:
        return ""
    return text.strip()


def extract_pdf_ocr(file_path: str) -> str:
    text = ""
    try:
        pdf = fitz.open(file_path)
        for page in pdf:
            try:
                pix = page.get_pixmap(dpi=300)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                extracted = pytesseract.image_to_string(img)
                if extracted:
                    text += extracted + "\n"
            except Exception:
                continue
    except Exception:
        return ""
    return text.strip()


def extract_text_from_pdf(file_path: str) -> str:
    try:
        text = extract_pdf_pypdf2(file_path)
        if len(text) > 300:
            return text

        text = extract_pdf_pymupdf(file_path)
        if len(text) > 300:
            return text

        return extract_pdf_ocr(file_path)
    except Exception:
        return ""


# ---------------- DOCX ---------------- #

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
        content = []

        for para in doc.paragraphs:
            if para.text and para.text.strip():
                content.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                try:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        content.append(row_text)
                except Exception:
                    continue

        return "\n".join(content).strip()

    except Exception:
        return ""


# ---------------- TXT ---------------- #

def extract_text_from_txt(file_path: str) -> str:
    encodings = ["utf-8", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc, errors="ignore") as f:
                return f.read().strip()
        except Exception:
            continue
    return ""


# ---------------- MAIN ---------------- #

def extract_text(file_path: str) -> str:
    try:
        if not file_path or not os.path.exists(file_path):
            return ""

        lower = file_path.lower()

        if lower.endswith(".pdf"):
            return extract_text_from_pdf(file_path)

        if lower.endswith(".docx"):
            return extract_text_from_docx(file_path)

        if lower.endswith(".txt"):
            return extract_text_from_txt(file_path)

        return ""

    except Exception:
        return ""
